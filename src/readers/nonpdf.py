import os
import uuid
import time

import cv2
import argparse
import requests
import numpy as np
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt
from operator import itemgetter
from functools import cmp_to_key
from sklearn.cluster import MeanShift
from docx.enum.style import WD_STYLE_TYPE
from pdf2image import convert_from_bytes

CHAR_PER_LINE = 159 #font size 13   line width 1300

load_dotenv()


def parse_args():
    parser = argparse.ArgumentParser(description="Convert pdf to docx")
    parser.add_argument("--fpath", type=str, help="Path to pdf folder")
    parser.add_argument("--fopath", type=str, help="Path to output docx folder")
    return parser.parse_args()


def line_clustering(line_list):
    y_list = []
    for line in line_list:
        y_list.append([line[5]])
    X = np.array(y_list)
    clustering = MeanShift(bandwidth=2).fit(X)
    
    line_dicts = {}
    for i, label in enumerate(clustering.labels_):
        if label in line_dicts.keys():
            line_dicts[label].append(line_list[i])
        else:
            line_dicts[label] = [line_list[i]]
            
    return line_dicts

def vietocr(roi):
    url = os.getenv("PDF_THUANN_URL")
    file_bytes = cv2.imencode('.jpg', roi)[1].tobytes()
    f = {'binary_file': file_bytes}
    response = requests.post(url, files = f)
    response = response.json()
    return response['predicts'][0][0]

def text_reg(imgs):
    VIETOCR_URL = os.getenv("VIETOCR_URL")
    f = []
    for img in imgs:
        try:
            file_bytes = cv2.imencode(".jpg", img)[1].tobytes()
        except:
            return [{"str": ""}]
        f.append(("binary_files", file_bytes))

    response = requests.post(VIETOCR_URL, files=f)
    response = response.json()
    # print(response)
    return response["predicts"]

def textline_detect(img):
    TEXTLINE_URL = (
        os.getenv("TEXTLINE_URL")
    )

    file_bytes = cv2.imencode(".jpg", img)[1].tobytes()
    f = {"binary_file": file_bytes}
    data = {"threshold": 0.5}

    response = requests.post(TEXTLINE_URL, files=f, data=data)
    response = response.json()

    return response["predicts"][0]

def convert_single(fpath, document):
    
    img = cv2.imread(fpath)
    imgh, imgw, _ = img.shape
    
    detect_result = textline_detect(img)
    
    line_list = []
    for box in detect_result:
        cls = box["cls"]
        bbox = box["2point"]
        x1 = bbox[0]
        y1 = bbox[1]
        x2 = bbox[2]
        y2 = bbox[3]
        cx = (x1+x2)/2
        cy = (y1+y2)/2
        line_list.append([x1, y1, x2, y2, cx, cy])
    
    line_dicts = line_clustering(line_list)
    
    line_group_list = x = [[] for i in range(max(line_dicts.keys())+1)]
    for cluster in line_dicts.keys():
        for line in line_dicts[cluster]:
            line_group_list[int(cluster)].append(line)
    
    line_group_list = sorted(line_group_list, key=cmp_to_key(lambda item1, item2: item1[0][5] - item2[0][5]))
    for i in range(len(line_group_list)):
        line_group_list[i] = sorted(line_group_list[i], key=itemgetter(4))
    
#    for cluster in line_group_list:
#        for line in cluster:
#            img = cv2.putText(img, str(line_group_list.index(cluster)), (int(line[4]), int(line[5])), cv2.FONT_HERSHEY_SIMPLEX, 1, (0,0,255), 2, cv2.LINE_AA)
#    cv2.imwrite("out.jpg", img)

    
    
    font_styles = document.styles
    comments_style = 'CommentsStyle_%s' %uuid.uuid4()
    font_charstyle = font_styles.add_style(comments_style, WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(13)
    font_object.name = 'Times New Roman'
    
    style = document.styles['Normal']
    style.paragraph_format.line_spacing = 1
    
#    
    for cluster in line_group_list:
        sentence = ""
        latest_roi_width = 0
        roi_list = []
        for i, line in enumerate(cluster):
            x1 = line[0]
            y1 = line[1]
            x2 = line[2]
            y2 = line[3]
            # print("line width ", x2-x1)
            roi = img[y1:y2, x1:x2]
            space_num = int(((x1-latest_roi_width)/imgw)*CHAR_PER_LINE)
            # print("space_num", space_num)
            if i == 0:
                tmp_sentence = " "*(space_num-22)
            else:
                tmp_sentence = " "*(space_num)

            roi_list.append(roi)

        reg_result_list = text_reg(roi_list)
        for reg_result in reg_result_list:
            text = reg_result["str"]
            tmp_sentence += "%s" %text
            sentence += tmp_sentence
            latest_roi_width = x2
        parag = document.add_paragraph()
        parag.style = 'Body Text'
        parag.add_run(sentence, style=comments_style)

def convert(fpath, file_out):
    imgpaths = []
    imgs = convert_from_bytes(open(fpath, "rb").read())
    for img in imgs:    
        fname = uuid.uuid4()
        fpath = 'data_imgs/%s.jpg' %fname
        if not os.path.exists('data_imgs'):
            os.makedirs('data_imgs')
        img.save(fpath, 'JPEG')
        imgpaths.append(fpath)
    
    document = Document()
    
    for imgpath in imgpaths:
        # print(imgpath)
        convert_single(imgpath, document)
        os.remove(imgpath)
    
    document.save(file_out)
    # return document
    
# if __name__ == "__main__":    
#     args = parse_args()
#     fpath = args.fpath
#     fopath = args.fopath
    
#     Path(fopath).mkdir(parents=True, exist_ok=True)
    
#     for file_pdf in Path(fpath).iterdir():
#         file_pdf = Path("data/475-qd-dhcntt_20-8-2018_qui_dinh_ve_tieu_chuan_giang_vien_giang_day_mon_hoc_va_tro_giang_mon_hoc.pdf")
#         file_out = Path(fopath) / (file_pdf.stem + ".docx")
#         print(file_pdf)
#         print(file_out)
#         start_time = time.time()
#         convert(str(file_pdf), str(file_out))
#         print(f"Convert {file_pdf} to {file_out} in {time.time()-start_time} seconds")
#         break
    
